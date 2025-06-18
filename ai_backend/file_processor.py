import os
import io
import logging
import hashlib
import magic
import re
from typing import Dict, Any, Optional

import PyPDF2
import pdfplumber
from docx import Document
import olefile
import chardet

logger = logging.getLogger(__name__)

class FileProcessor:
    """파일 분석 및 텍스트 추출 서비스 (보안 강화)"""
    
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'PDF 문서',
        '.docx': 'Word 문서',
        '.txt': '텍스트 파일',
    }
    
    # MIME 타입 검증
    ALLOWED_MIME_TYPES = {
        '.pdf': ['application/pdf'],
        '.docx': [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ],
        '.txt': [
            'text/plain',
            'text/x-python',
            'text/x-script.python',
            'application/octet-stream'  # 일부 텍스트 파일이 이렇게 감지됨
        ]
    }
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    MAX_TEXT_LENGTH = 1000000  # 1MB 텍스트
    
    # 위험한 패턴들
    DANGEROUS_PATTERNS = [
        r'<script[^>]*>.*?</script>',
        r'javascript:',
        r'vbscript:',
        r'on\w+\s*=',
        r'\\x[0-9a-f]{2}',  # hex encoded
        r'%[0-9a-f]{2}',    # url encoded
    ]
    
    @classmethod
    def is_supported_file(cls, filename: str) -> bool:
        """지원되는 파일 형식인지 확인"""
        _, ext = os.path.splitext(cls.sanitize_filename(filename).lower())
        return ext in cls.SUPPORTED_EXTENSIONS
    
    @classmethod
    def sanitize_filename(cls, filename: str) -> str:
        """파일명 sanitization"""
        # 위험한 문자 제거
        filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        # 연속된 점 제거 (디렉토리 트래버설 방지)
        filename = re.sub(r'\.{2,}', '.', filename)
        # 길이 제한
        if len(filename) > 255:
            name, ext = os.path.splitext(filename)
            filename = name[:255-len(ext)] + ext
        return filename
    
    @classmethod
    def validate_file_security(cls, file_content: bytes, filename: str) -> Dict[str, Any]:
        """파일 보안 검증"""
        try:
            # 1. 파일명 검증
            safe_filename = cls.sanitize_filename(filename)
            _, ext = os.path.splitext(safe_filename.lower())
            
            if ext not in cls.SUPPORTED_EXTENSIONS:
                return {
                    'valid': False,
                    'error': f'지원되지 않는 파일 형식: {ext}'
                }
            
            # 2. 파일 크기 검증
            if len(file_content) > cls.MAX_FILE_SIZE:
                return {
                    'valid': False,
                    'error': f'파일 크기가 너무 큽니다 (최대 {cls.MAX_FILE_SIZE // (1024*1024)}MB)'
                }
            
            # 3. 빈 파일 검증
            if len(file_content) == 0:
                return {
                    'valid': False,
                    'error': '빈 파일입니다'
                }
            
            # 4. MIME 타입 검증 (python-magic 사용)
            try:
                mime_type = magic.from_buffer(file_content, mime=True)
                allowed_mimes = cls.ALLOWED_MIME_TYPES.get(ext, [])
                
                if mime_type not in allowed_mimes:
                    return {
                        'valid': False,
                        'error': f'파일 내용이 확장자와 일치하지 않습니다. 감지된 타입: {mime_type}'
                    }
            except Exception as e:
                logger.warning(f"MIME 타입 검증 실패: {e}")
                # MIME 검증 실패시 파일 헤더로 기본 검증
                if not cls._validate_file_header(file_content, ext):
                    return {
                        'valid': False,
                        'error': '파일 형식이 올바르지 않습니다'
                    }
            
            # 5. 파일 헤더 검증
            if not cls._validate_file_header(file_content, ext):
                return {
                    'valid': False,
                    'error': '파일 헤더가 올바르지 않습니다'
                }
            
            # 6. 파일 해시 생성 (중복 방지 및 추적용)
            file_hash = hashlib.sha256(file_content).hexdigest()
            
            return {
                'valid': True,
                'safe_filename': safe_filename,
                'file_hash': file_hash,
                'mime_type': mime_type if 'mime_type' in locals() else 'unknown'
            }
            
        except Exception as e:
            logger.error(f"파일 보안 검증 중 오류: {e}")
            return {
                'valid': False,
                'error': '파일 검증 중 오류가 발생했습니다'
            }
    
    @classmethod
    def _validate_file_header(cls, file_content: bytes, ext: str) -> bool:
        """파일 헤더 검증"""
        if len(file_content) < 4:
            return False
            
        header = file_content[:8]
        
        if ext == '.pdf':
            return header.startswith(b'%PDF-')
        elif ext == '.docx':
            # DOCX는 ZIP 파일이므로 ZIP 헤더 확인
            return header.startswith(b'PK\x03\x04') or header.startswith(b'PK\x05\x06')
        elif ext == '.txt':
            # 텍스트 파일은 UTF-8, ASCII 등으로 디코딩 가능한지 확인
            try:
                detected = chardet.detect(file_content[:1024])
                return detected.get('confidence', 0) > 0.5
            except:
                return True  # 텍스트는 관대하게 처리
        
        return False
    
    @classmethod
    def sanitize_text_content(cls, text: str) -> str:
        """텍스트 내용 sanitization"""
        if len(text) > cls.MAX_TEXT_LENGTH:
            text = text[:cls.MAX_TEXT_LENGTH] + "\n\n[텍스트가 너무 길어 일부가 잘렸습니다]"
        
        # 위험한 패턴 제거
        for pattern in cls.DANGEROUS_PATTERNS:
            text = re.sub(pattern, '[제거된 내용]', text, flags=re.IGNORECASE)
        
        # 제어 문자 제거 (인쇄 가능한 문자와 일반적인 공백문자만 허용)
        text = ''.join(char for char in text if char.isprintable() or char in '\n\r\t ')
        
        return text
    
    @classmethod
    def is_supported_file(cls, filename: str) -> bool:
        """지원되는 파일 형식인지 확인"""
        _, ext = os.path.splitext(filename.lower())
        return ext in cls.SUPPORTED_EXTENSIONS
    
    @classmethod
    def get_file_type(cls, filename: str) -> str:
        """파일 타입 반환"""
        _, ext = os.path.splitext(filename.lower())
        return cls.SUPPORTED_EXTENSIONS.get(ext, '알 수 없는 형식')
    
    @classmethod
    def extract_text_from_file(cls, file_content: bytes, filename: str) -> Dict[str, Any]:
        """파일에서 텍스트를 추출합니다 (보안 검증 포함)"""
        try:
            # 1. 보안 검증 먼저 수행
            security_check = cls.validate_file_security(file_content, filename)
            if not security_check['valid']:
                return {
                    'success': False,
                    'error': security_check['error'],
                    'text': '',
                    'metadata': {'security_error': True}
                }
            
            safe_filename = security_check['safe_filename']
            file_hash = security_check['file_hash']
            
            # 2. 파일 확장자별 처리
            _, ext = os.path.splitext(safe_filename.lower())
            
            if ext == '.pdf':
                result = cls._extract_pdf_text(file_content, safe_filename)
            elif ext == '.docx':
                result = cls._extract_docx_text(file_content, safe_filename)
            elif ext == '.txt':
                result = cls._extract_txt_text(file_content, safe_filename)
            else:
                return {
                    'success': False,
                    'error': f'지원되지 않는 파일 형식: {ext}',
                    'text': '',
                    'metadata': {}
                }
            
            # 3. 추출된 텍스트 sanitization
            if result['success'] and result['text']:
                result['text'] = cls.sanitize_text_content(result['text'])
                result['metadata']['file_hash'] = file_hash
                result['metadata']['safe_filename'] = safe_filename
                result['metadata']['original_filename'] = filename
                result['metadata']['security_validated'] = True
            
            return result
                
        except Exception as e:
            logger.error(f"파일 처리 중 오류 발생: {e}")
            return {
                'success': False,
                'error': '파일 처리 중 보안 오류가 발생했습니다',
                'text': '',
                'metadata': {'security_error': True}
            }
    
    @classmethod
    def _extract_pdf_text(cls, file_content: bytes, filename: str) -> Dict[str, Any]:
        """PDF 파일에서 텍스트 추출"""
        try:
            # pdfplumber를 먼저 시도 (더 정확한 텍스트 추출)
            try:
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    text_parts = []
                    metadata = {
                        'pages': len(pdf.pages),
                        'file_type': 'PDF',
                        'extraction_method': 'pdfplumber'
                    }
                    
                    for page_num, page in enumerate(pdf.pages[:50], 1):  # 최대 50페이지까지
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(f"=== 페이지 {page_num} ===\n{page_text}\n")
                        except Exception as e:
                            text_parts.append(f"=== 페이지 {page_num} ===\n[텍스트 추출 실패: {str(e)}]\n")
                    
                    text = '\n'.join(text_parts)
                    
                    if not text.strip():
                        raise Exception("텍스트를 추출할 수 없습니다")
                    
                    return {
                        'success': True,
                        'text': text,
                        'metadata': metadata,
                        'error': None
                    }
            
            except Exception:
                # pdfplumber 실패시 PyPDF2로 시도
                reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text_parts = []
                metadata = {
                    'pages': len(reader.pages),
                    'file_type': 'PDF',
                    'extraction_method': 'PyPDF2'
                }
                
                for page_num, page in enumerate(reader.pages[:50], 1):  # 최대 50페이지까지
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"=== 페이지 {page_num} ===\n{page_text}\n")
                    except Exception as e:
                        text_parts.append(f"=== 페이지 {page_num} ===\n[텍스트 추출 실패: {str(e)}]\n")
                
                text = '\n'.join(text_parts)
                
                if not text.strip():
                    return {
                        'success': False,
                        'error': 'PDF에서 텍스트를 추출할 수 없습니다. 이미지 기반 PDF이거나 보호된 문서일 수 있습니다.',
                        'text': '',
                        'metadata': metadata
                    }
                
                return {
                    'success': True,
                    'text': text,
                    'metadata': metadata,
                    'error': None
                }
        
        except Exception as e:
            return {
                'success': False,
                'error': f'PDF 처리 중 오류: {str(e)}',
                'text': '',
                'metadata': {'file_type': 'PDF'}
            }
    
    @classmethod
    def _extract_docx_text(cls, file_content: bytes, filename: str) -> Dict[str, Any]:
        """DOCX 파일에서 텍스트 추출"""
        try:
            doc = Document(io.BytesIO(file_content))
            
            text_parts = []
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    paragraph_count += 1
            
            # 표(table) 내용도 추출
            table_count = 0
            for table in doc.tables:
                table_count += 1
                text_parts.append(f"\n=== 표 {table_count} ===")
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            text = '\n'.join(text_parts)
            
            metadata = {
                'file_type': 'DOCX',
                'paragraphs': paragraph_count,
                'tables': table_count,
                'extraction_method': 'python-docx'
            }
            
            if not text.strip():
                return {
                    'success': False,
                    'error': 'DOCX 파일에서 텍스트를 찾을 수 없습니다.',
                    'text': '',
                    'metadata': metadata
                }
            
            return {
                'success': True,
                'text': text,
                'metadata': metadata,
                'error': None
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': f'DOCX 처리 중 오류: {str(e)}',
                'text': '',
                'metadata': {'file_type': 'DOCX'}
            }
    
    @classmethod
    def _extract_txt_text(cls, file_content: bytes, filename: str) -> Dict[str, Any]:
        """TXT 파일에서 텍스트 추출"""
        try:
            # 인코딩 자동 감지
            detected = chardet.detect(file_content)
            encoding = detected.get('encoding', 'utf-8')
            confidence = detected.get('confidence', 0)
            
            # 인코딩 시도 순서
            encodings_to_try = [encoding, 'utf-8', 'cp949', 'euc-kr', 'latin1']
            
            text = None
            used_encoding = None
            
            for enc in encodings_to_try:
                if enc:
                    try:
                        text = file_content.decode(enc)
                        used_encoding = enc
                        break
                    except (UnicodeDecodeError, LookupError):
                        continue
            
            if text is None:
                return {
                    'success': False,
                    'error': '텍스트 파일의 인코딩을 해석할 수 없습니다.',
                    'text': '',
                    'metadata': {'file_type': 'TXT'}
                }
            
            metadata = {
                'file_type': 'TXT',
                'encoding': used_encoding,
                'detection_confidence': confidence,
                'lines': len(text.splitlines()),
                'extraction_method': 'chardet'
            }
            
            return {
                'success': True,
                'text': text,
                'metadata': metadata,
                'error': None
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': f'TXT 처리 중 오류: {str(e)}',
                'text': '',
                'metadata': {'file_type': 'TXT'}
            }
    
    @classmethod
    def get_file_summary(cls, text: str, metadata: Dict[str, Any]) -> str:
        """파일 정보 요약 생성"""
        file_type = metadata.get('file_type', '알 수 없음')
        summary_parts = [f"파일 형식: {file_type}"]
        
        if 'pages' in metadata:
            summary_parts.append(f"페이지 수: {metadata['pages']}")
        
        if 'paragraphs' in metadata:
            summary_parts.append(f"문단 수: {metadata['paragraphs']}")
        
        if 'tables' in metadata:
            summary_parts.append(f"표 개수: {metadata['tables']}")
        
        if 'lines' in metadata:
            summary_parts.append(f"줄 수: {metadata['lines']}")
        
        if 'encoding' in metadata:
            summary_parts.append(f"인코딩: {metadata['encoding']}")
        
        text_length = len(text)
        if text_length > 0:
            summary_parts.append(f"텍스트 길이: {text_length:,}자")
        
        if metadata.get('security_validated'):
            summary_parts.append("보안 검증 완료")
        
        return " | ".join(summary_parts)
